#!/usr/bin/env python3
"""
착공신고서 템플릿 생성 스크립트
Clean DOCX template generator for construction reports
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_border(cell, **kwargs):
    """Add borders to table cell"""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()

    # Create borders element
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_element = OxmlElement(f'w:{edge}')
        edge_element.set(qn('w:val'), 'single')
        edge_element.set(qn('w:sz'), '4')
        edge_element.set(qn('w:space'), '0')
        edge_element.set(qn('w:color'), '000000')
        tcBorders.append(edge_element)

    tcPr.append(tcBorders)

def create_template():
    """Create clean construction report template"""
    doc = Document()

    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('사물인터넷(IoT) 측정기기 부착지원 사업\n착 공 신 고 서')
    title_run.font.size = Pt(18)
    title_run.font.bold = True

    doc.add_paragraph()  # Spacing

    # Main information table
    table = doc.add_table(rows=11, cols=2)
    table.style = 'Table Grid'

    # Set column widths
    table.columns[0].width = Inches(1.5)
    table.columns[1].width = Inches(5.0)

    # Row 1: 사업장명
    cell = table.rows[0].cells[0]
    cell.text = '사 업 장 명'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.paragraphs[0].runs[0].font.bold = True
    add_border(cell)

    cell = table.rows[0].cells[1]
    cell.text = '{{사업장명}}'
    add_border(cell)

    # Row 2: 사업장 소재지 (주소)
    cell = table.rows[1].cells[0]
    cell.text = '사업장소재지'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.paragraphs[0].runs[0].font.bold = True
    add_border(cell)

    cell = table.rows[1].cells[1]
    p = cell.paragraphs[0]
    p.add_run('{{주소}}\n')
    p.add_run('전화: {{회사연락처}}  팩스: {{팩스번호}}')
    add_border(cell)

    # Row 3: 설치업체명
    cell = table.rows[2].cells[0]
    cell.text = '설 치 업 체 명'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.paragraphs[0].runs[0].font.bold = True
    add_border(cell)

    cell = table.rows[2].cells[1]
    cell.text = '주식회사 블루온'
    add_border(cell)

    # Row 4: 시공업체 소재지
    cell = table.rows[3].cells[0]
    cell.text = '시공업체소재지'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.paragraphs[0].runs[0].font.bold = True
    add_border(cell)

    cell = table.rows[3].cells[1]
    p = cell.paragraphs[0]
    p.add_run('경상북도 고령군 대가야읍 낫질로 285\n')
    p.add_run('전화: 1661-5543  팩스: 031-8077-2054')
    add_border(cell)

    # Row 5: 사업자등록번호
    cell = table.rows[4].cells[0]
    cell.text = '사업자등록번호'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.paragraphs[0].runs[0].font.bold = True
    add_border(cell)

    cell = table.rows[4].cells[1]
    p = cell.paragraphs[0]
    p.add_run('사업장: {{사업자등록번호}}\n')
    p.add_run('시공업체: 679-86-02827')
    add_border(cell)

    # Row 6: 부착기간
    cell = table.rows[5].cells[0]
    cell.text = '부 착 기 간'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.paragraphs[0].runs[0].font.bold = True
    add_border(cell)

    cell = table.rows[5].cells[1]
    cell.text = '{{보조금 승인일}} 부터 {{보조금 승인일+3개월}} 까지 (3개월)'
    add_border(cell)

    # Row 7: 총 소요금액
    cell = table.rows[6].cells[0]
    cell.text = '총 소 요 금 액'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.paragraphs[0].runs[0].font.bold = True
    add_border(cell)

    cell = table.rows[6].cells[1]
    cell.text = '{{환경부고시가}} 원'
    add_border(cell)

    # Row 8: 보조금 승인액
    cell = table.rows[7].cells[0]
    cell.text = '보조금승인액'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.paragraphs[0].runs[0].font.bold = True
    add_border(cell)

    cell = table.rows[7].cells[1]
    cell.text = '{{보조금 승인액}} 원'
    add_border(cell)

    # Row 9: 자체부담액
    cell = table.rows[8].cells[0]
    cell.text = '자 체 부 담 액'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.paragraphs[0].runs[0].font.bold = True
    add_border(cell)

    cell = table.rows[8].cells[1]
    cell.text = '{{자부담}} 원'
    add_border(cell)

    # Row 10: 입금액
    cell = table.rows[9].cells[0]
    cell.text = '입 금 액'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.paragraphs[0].runs[0].font.bold = True
    add_border(cell)

    cell = table.rows[9].cells[1]
    cell.text = '{{입금액}} 원 (부가세 포함)'
    add_border(cell)

    # Row 11: 설치 품목
    cell = table.rows[10].cells[0]
    cell.text = '설 치 품 목'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.paragraphs[0].runs[0].font.bold = True
    add_border(cell)

    cell = table.rows[10].cells[1]
    p = cell.paragraphs[0]
    p.add_run('게이트웨이: {{게이트웨이}}대, VPN: {{VPN}}, ')
    p.add_run('배출CT: {{배출CT}}개, 방지CT: {{방지CT}}개, ')
    p.add_run('차압계: {{차압계}}개, 온도계: {{온도계}}개, PH계: {{PH계}}개\n')
    p.add_run('방지시설: {{방지시설명}}')
    add_border(cell)

    doc.add_paragraph()  # Spacing

    # Declaration text
    declaration = doc.add_paragraph()
    declaration.alignment = WD_ALIGN_PARAGRAPH.CENTER
    declaration_run = declaration.add_run(
        '소규모 사업장 사물인터넷(IoT) 측정기기 부착지원 사업에\n'
        '대하여 착공신고서를 제출합니다.'
    )
    declaration_run.font.size = Pt(11)

    doc.add_paragraph()  # Spacing

    # Date and signature
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run('{{year}} 년  {{month}} 월  {{day}} 일')
    date_run.font.size = Pt(12)

    doc.add_paragraph()  # Spacing

    # Applicant signature
    signature = doc.add_paragraph()
    signature.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sig_run = signature.add_run('신청인(대표자)  {{사업장명}}  {{대표자성명}}  (인)')
    sig_run.font.size = Pt(11)

    doc.add_paragraph()  # Spacing

    # Recipient
    recipient = doc.add_paragraph()
    recipient.alignment = WD_ALIGN_PARAGRAPH.CENTER
    recipient_run = recipient.add_run('{{지자체장}}  귀하')
    recipient_run.font.size = Pt(12)
    recipient_run.font.bold = True

    doc.add_paragraph()  # Spacing

    # Required documents table
    docs_table = doc.add_table(rows=1, cols=1)
    docs_table.style = 'Table Grid'
    cell = docs_table.rows[0].cells[0]

    header = cell.paragraphs[0]
    header.add_run('구 비 서 류').font.bold = True
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = cell.add_paragraph()
    p.add_run('1. 대기배출시설 설치 허가(신고)증 사본 1부.\n')
    p.add_run('2. 계약서(사본) 1부.\n')
    p.add_run('3. 자부담금 입금 확인증 1부.\n')
    p.add_run('4. 계약이행보증보험 1부.\n')
    p.add_run('5. 개선계획서(최종, 보완사항 포함) 1부.')

    add_border(cell)

    # Save the document
    output_path = '양식/☆착공신고서 템플릿_깨끗한버전.docx'
    doc.save(output_path)
    print(f'✅ 템플릿 파일이 생성되었습니다: {output_path}')
    print('\n📋 포함된 플레이스홀더:')
    placeholders = [
        '사업장명', '주소', '회사연락처', '팩스번호', '사업자등록번호',
        '보조금 승인일', '보조금 승인일+3개월', '환경부고시가', '보조금 승인액',
        '자부담', '입금액', '게이트웨이', 'VPN', '배출CT', '방지CT',
        '차압계', '온도계', 'PH계', '방지시설명', 'year', 'month', 'day',
        '대표자성명', '지자체장'
    ]
    for i, placeholder in enumerate(placeholders, 1):
        print(f'  {i}. {{{{{placeholder}}}}}')

if __name__ == '__main__':
    create_template()
